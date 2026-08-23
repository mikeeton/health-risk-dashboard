from pathlib import Path
import json, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "dissertation" / "dissertation_source.md"
OUT = ROOT / "docs" / "dissertation" / "Health_Risk_Dashboard_Dissertation.docx"

def field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText"); text.set(qn("xml:space"), "preserve"); text.text = instruction
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    val = OxmlElement("w:t"); val.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for node in (begin, text, sep, val, end): run._r.append(node)

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def repeat_header(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement("w:tblHeader"); el.set(qn("w:val"), "true"); trPr.append(el)

def page_number(footer):
    p = footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; field(p, "PAGE")

def set_page_number_format(section, fmt, start=1):
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None: pg = OxmlElement("w:pgNumType"); sectPr.append(pg)
    pg.set(qn("w:fmt"), fmt); pg.set(qn("w:start"), str(start))

def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.italic = True

def add_static_toc(doc):
    entries = [
        ("Chapter One: Introduction", "1"),
        ("Chapter Two: Literature Review", "5"),
        ("Chapter Three: Requirements Analysis and Methodology", "13"),
        ("Chapter Four: System Analysis and Design", "18"),
        ("Chapter Five: System Implementation", "26"),
        ("Chapter Six: Testing and Evaluation", "39"),
        ("Chapter Seven: Results and Discussion", "45"),
        ("Chapter Eight: Conclusion and Recommendations", "47"),
        ("References", "49"),
        ("Appendix A: Reproducibility Checklist", "52"),
    ]
    for label, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(5.9), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        run = p.add_run(f"{label}\t{page}")
        run.bold = label.startswith("Chapter")

def add_static_list(doc, entries):
    for number, label in entries:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(4)
        number_run = p.add_run(f"{number}  ")
        number_run.bold = True
        p.add_run(label)

def add_lists_of_illustrations(doc):
    add_static_list(doc, [
        ("Figure 4.1", "Layered architecture of the Health Risk Dashboard"),
        ("Figure 4.2", "Role and patient-access boundary"),
        ("Figure 4.3", "Observation-to-alert data flow"),
        ("Figure 4.4", "Evidence-bound AI assistant pipeline"),
        ("Figure 4.5", "Production deployment topology"),
        ("Figure 5.1", "Doctor dashboard displaying an assigned patient and current observations"),
        ("Figure 5.2", "Nurse dashboard organised around bedside care activities"),
        ("Figure 5.3", "Administrator dashboard for non-clinical platform governance"),
        ("Figure 5.4", "Notification centre with unread inbox and retained history"),
        ("Figure 5.5", "Reproducible machine-learning lifecycle"),
        ("Figure 5.6", "Trained six-hour model output with probability and provenance"),
        ("Figure 5.7", "Restricted research evidence workspace"),
        ("Figure 5.8", "Evidence-linked clinician Groq AI Assistant"),
        ("Figure 5.9", "Responsive authentication interface"),
        ("Figure 6.1", "Candidate-model performance on the validation partition"),
        ("Figure 6.2", "Internal test confusion matrix at the 0.014 threshold"),
        ("Figure 6.3", "Internal test calibration plot"),
        ("Figure 6.4", "Internal test and external Set B metric comparison"),
        ("Figure 6.5", "Global SHAP feature influence in the evaluation sample"),
    ])
    doc.add_heading("List of Tables", level=1)
    add_static_list(doc, [
        ("Table 3.1", "Principal functional requirements"),
        ("Table 4.1", "Logical layers of the Health Risk Dashboard"),
        ("Table 4.2", "Summary threat model"),
        ("Table 6.1", "Final engineering verification summary"),
        ("Table 6.2", "Retrospective model performance at the selected operating threshold"),
        ("Table 6.3", "Evidence maturity and permitted conclusions"),
    ])

def add_figure(doc, filename, caption, alt_text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    image_path = ROOT / "docs" / "dissertation" / (filename if "/" in filename else f"figures/{filename}")
    # Keep the architecture figure slightly shorter so its paired summary table
    # remains together on the same page; other figures use the widest readable size.
    figure_width = 4.90 if filename == "architecture.png" else 5.72
    run = p.add_run(); run.add_picture(str(image_path), width=Inches(figure_width))
    drawing = run._r.find(qn("w:drawing"))
    if drawing is not None:
        docPr = drawing.find(".//" + qn("wp:docPr"))
        if docPr is not None: docPr.set("descr", alt_text); docPr.set("title", caption)
    add_caption(doc, caption + " Source: author-generated from the implemented system and evaluation artifact.")

def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = h; set_cell_shading(c, "E7E6E6"); c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraphs[0].paragraph_format.first_line_indent = Inches(0)
        for run in c.paragraphs[0].runs: run.bold = True
    repeat_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            cells[i].paragraphs[0].paragraph_format.first_line_indent = Inches(0)
    # A4 with a 1.5-inch binding margin and 1-inch right margin leaves
    # approximately 8,310 twips.  Write consistent tblW/grid/tcW geometry so
    # Word and LibreOffice do not reinterpret an auto-width table differently.
    weights = widths or [1] * len(headers)
    raw = [round(8310 * float(w) / sum(weights)) for w in weights]
    raw[-1] += 8310 - sum(raw)
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None: tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa"); tblW.set(qn("w:w"), "8310")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for value in raw:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(value)); grid.append(col)
    for row in table.rows:
        for i, value in enumerate(raw):
            tcW = row.cells[i]._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcW is None: tcW = OxmlElement("w:tcW"); row.cells[i]._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn("w:type"), "dxa"); tcW.set(qn("w:w"), str(value))
    return table

def add_architecture(doc):
    add_table(doc, ["Layer", "Principal components", "Responsibility"], [
        ["Presentation", "React, TypeScript, Vite, role-aware routes", "Accessible dashboards, workflows and evidence display"],
        ["Application", "FastAPI, Pydantic, JWT, access-control dependencies", "Validation, authorisation, clinical workflows and APIs"],
        ["Data and messaging", "PostgreSQL, Alembic, Redis, WebSockets", "Durable records, migrations and live notifications"],
        ["Intelligence", "Deterministic rules, calibrated ML, SHAP, Groq", "Safety escalation, prediction and evidence-linked assistance"],
        ["Research", "Shadow validation, SUS, effectiveness records", "Restricted evaluation without altering clinical records"],
    ], [1.1, 2.55, 2.85])
    add_caption(doc, "Table 4.1: Logical layers of the Health Risk Dashboard.")

def add_metrics(doc):
    e = json.loads((ROOT / "backend" / "artifacts" / "ml" / "evaluation.json").read_text())
    m = e["test_metrics"]; x = e["external_validation"]["metrics"]
    add_table(doc, ["Measure", "Internal test", "External Set B"], [
        ["ROC-AUC", f'{m["roc_auc"]:.3f}', f'{x["roc_auc"]:.3f}'],
        ["PR-AUC", f'{m["pr_auc"]:.3f}', f'{x["pr_auc"]:.3f}'],
        ["Sensitivity", f'{m["recall_sensitivity"]:.3f}', f'{x["recall_sensitivity"]:.3f}'],
        ["Specificity", f'{m["specificity"]:.3f}', f'{x["specificity"]:.3f}'],
        ["Precision", f'{m["precision"]:.3f}', f'{x["precision"]:.3f}'],
        ["Brier score", f'{m["brier_score"]:.3f}', f'{x["brier_score"]:.3f}'],
    ], [2.5, 2, 2])
    add_caption(doc, "Table 6.2: Retrospective model performance at the selected operating threshold.")

def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]; normal.font.name = "Times New Roman"; normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format; pf.line_spacing = 1.5; pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; pf.first_line_indent = Inches(.5); pf.space_after = Pt(0)
    for name, size in [("Title", 20), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        s = styles[name]; s.font.name = "Times New Roman"; s.font.size = Pt(size); s.font.bold = True; s.font.color.rgb = RGBColor(0,0,0)
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        s.paragraph_format.keep_with_next = True
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        s.paragraph_format.first_line_indent = Inches(0)
    styles["Heading 1"].paragraph_format.page_break_before = True
    styles["Caption"].font.name = "Times New Roman"; styles["Caption"].font.size = Pt(10); styles["Caption"].font.color.rgb = RGBColor(0,0,0)

def parse(doc, text):
    lines = text.splitlines(); i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line: i += 1; continue
        if line == "[[ARCHITECTURE_TABLE]]": add_architecture(doc); i += 1; continue
        if line == "[[METRICS_TABLE]]": add_metrics(doc); i += 1; continue
        if line.startswith("[[FIGURE|"):
            parts=line[2:-2].split("|",3); add_figure(doc,parts[1],parts[2],parts[3]); i += 1; continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1); i += 1; continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2); i += 1; continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3); i += 1; continue
        if line.startswith("- "):
            while i < len(lines) and lines[i].startswith("- "):
                p=doc.add_paragraph(style="List Bullet"); p.add_run(lines[i][2:]); i += 1
            continue
        if line.startswith("TABLE|"):
            block=[]
            while i < len(lines) and lines[i].startswith("TABLE|"):
                block.append(lines[i][6:].split("|")); i += 1
            add_table(doc, block[0], block[1:]); continue
        if line.startswith("CAPTION|"):
            add_caption(doc, line[8:]); i += 1; continue
        para=[line]; i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r'^(#{1,3} |-|TABLE\||CAPTION\||\[\[)', lines[i]):
            para.append(lines[i].strip()); i += 1
        p=doc.add_paragraph(" ".join(para)); p.paragraph_format.widow_control=True

def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc=Document(); configure_styles(doc)
    sec=doc.sections[0]; sec.page_width=Inches(8.27); sec.page_height=Inches(11.69); sec.left_margin=Inches(1.5); sec.right_margin=Inches(1); sec.top_margin=Inches(1); sec.bottom_margin=Inches(1)
    # Cover
    for _ in range(4): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("DESIGN AND IMPLEMENTATION OF AN AI-ASSISTED HEALTH RISK MONITORING DASHBOARD"); r.bold=True; r.font.size=Pt(18)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("A Full-Stack Clinical Monitoring and Research Prototype").italic=True
    for _ in range(3): doc.add_paragraph()
    for line in ["Michael Etonyeaku", "A dissertation submitted in partial fulfilment of the requirements for the award of a degree", "August 2026"]:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(line)
    doc.add_page_break()
    doc.add_heading("Declaration", level=1); doc.add_paragraph("This dissertation reports the design, implementation and evaluation of the Health Risk Dashboard. Sources are acknowledged using Harvard referencing. Software-generated evidence is distinguished from clinical evidence, and no claim of regulatory approval or demonstrated clinical effectiveness is made.")
    doc.add_heading("Acknowledgements", level=1); doc.add_paragraph("The project benefited from publicly available research, open-source software, the PhysioNet/CinC Challenge 2012 dataset and the guidance of academic and technical reviewers. Appreciation is extended to those whose feedback informed the iterative refinement of the interface, security controls and evaluation plan.")
    doc.add_heading("Abstract", level=1)
    abstract=(SOURCE.read_text(encoding="utf-8").split("<!--ABSTRACT-->")[1].split("<!--BODY-->")[0].strip())
    abstract_p = doc.add_paragraph(abstract)
    abstract_p.paragraph_format.line_spacing = 1.3
    for run in abstract_p.runs: run.font.size = Pt(11.5)
    doc.add_heading("Table of Contents", level=1); add_static_toc(doc)
    doc.add_heading("List of Figures", level=1); add_lists_of_illustrations(doc)
    doc.add_heading("List of Abbreviations", level=1)
    add_table(doc,["Abbreviation","Meaning"],[["AI","Artificial intelligence"],["API","Application programming interface"],["CSP","Content Security Policy"],["EHR","Electronic health record"],["JWT","JSON Web Token"],["ML","Machine learning"],["RBAC","Role-based access control"],["SHAP","SHapley Additive exPlanations"],["SUS","System Usability Scale"],["WCAG","Web Content Accessibility Guidelines"]],[1.8,4.7])
    set_page_number_format(sec, "lowerRoman", 1); page_number(sec.footer)
    main_sec=doc.add_section(WD_SECTION.NEW_PAGE); main_sec.left_margin=Inches(1.5); main_sec.right_margin=Inches(1); main_sec.top_margin=Inches(1); main_sec.bottom_margin=Inches(1); main_sec.header.is_linked_to_previous=False; main_sec.footer.is_linked_to_previous=False; set_page_number_format(main_sec,"decimal",1); page_number(main_sec.footer)
    hp=main_sec.header.paragraphs[0]; hp.text="AI-Assisted Health Risk Monitoring Dashboard"; hp.alignment=WD_ALIGN_PARAGRAPH.CENTER; hp.runs[0].font.size=Pt(9)
    body=SOURCE.read_text(encoding="utf-8").split("<!--BODY-->")[1]
    parse(doc, body)
    doc.core_properties.title="Design and Implementation of an AI-Assisted Health Risk Monitoring Dashboard"; doc.core_properties.author="Michael Etonyeaku"; doc.core_properties.subject="Final-year project dissertation"
    doc.save(OUT)
    words=len(re.findall(r"\b[\w'-]+\b", SOURCE.read_text(encoding="utf-8")))
    print(f"Saved {OUT}\nSource word count: {words}")

if __name__ == "__main__": main()
